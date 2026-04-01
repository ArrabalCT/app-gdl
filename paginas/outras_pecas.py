import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import io
import os
import re
from PIL import Image, ImageOps

# --- Funções de Formatação do Word ---
def adicionar_borda_inferior(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def adicionar_campo_numpages(paragraph):
    p = paragraph._p
    for char_type, text in [('begin', None), (None, ' NUMPAGES '), ('separate', None), (None, 'xx'), ('end', None)]:
        r = OxmlElement('w:r')
        if char_type:
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), char_type)
            r.append(fldChar)
        else:
            instrText = OxmlElement('w:instrText' if 'NUMPAGES' in text else 'w:t')
            if 'NUMPAGES' in text: instrText.set(qn('xml:space'), 'preserve')
            instrText.text = text
            r.append(instrText)
        p.append(r)

# --- Variáveis de Sessão ---
if 'itens_pecas' not in st.session_state: st.session_state['itens_pecas'] = []
if 'fotos' not in st.session_state: st.session_state['fotos'] = []
if 'fotos_removidas' not in st.session_state: st.session_state['fotos_removidas'] = []
if 'mk' not in st.session_state: st.session_state['mk'] = 0 
if 'uploader_key' not in st.session_state: st.session_state['uploader_key'] = 0
if 'item_mk' not in st.session_state: st.session_state['item_mk'] = 0 
mk = st.session_state['mk']
ik = st.session_state['item_mk']

# --- Listas de Dados ---
delegados = ["", "Adilson Antonio Marcondes dos Santos", "Adriane Goncalves", "Anisio Galdioli", "Cesar Aparecido Vieira da Silva", "Daniel Souza Baptista de Castro", "Ernani Ronaldo Giannico Braga", "Fabio Germano Figueiredo Cabett", "Flavia Maria Rocha Rollo", "Francisco Sannini Neto", "Hugo Parreiras de Macedo", "Jose Marcelo Silva Hial", "Leonardo da Costa Ferreira", "Marcelo Vieira Cavalcante", "Mario Celso Ribeiro Senne", "Paulo Roberto Gruschka Castilho", "Paulo Sergio Barbosa", "Pedro Rossati", "Sergio Lucas Adler Guedes de Oliveira", "Vania Idalira Z. de Oliveira", "Outro..."]
peritos = ["Alexandre Rabello de Oliveira", "Bruna Fernandes Nogueira", "Claude Thiago Arrabal", "Jéssica Pereira Gonçalves", "Júlia Soares Melo", "Luiz Fausto Prado Vasques", "Marcelo Mourão Dantas", "Márcio Steinmetz Soares", "Sarah Costa Teixeira", "Ruan Carvalho de Souza"]
cidades = ["", "Aparecida", "Cachoeira Paulista", "Canas", "Cunha", "Guaratinguetá", "Lorena", "Piquete", "Potim", "Roseira", "Outra..."]
dps_por_cidade = {
    "": [""],
    "Aparecida": ["", "DEL.POL.APARECIDA", "01º D.P. APARECIDA"],
    "Canas": ["", "DEL.POL.CANAS"],
    "Cachoeira Paulista": ["", "DEL.POL.CACHOEIRA PAULISTA", "01º D.P. CACHOEIRA PAULISTA"],
    "Cunha": ["", "DEL.POL.CUNHA"],
    "Guaratinguetá": ["", "01º D.P. GUARATINGUETA", "02º D.P. GUARATINGUETA", "03º D.P. GUARATINGUETA", "DEL.SEC.GUARATINGUETA PLANTÃO", "DISE- DEL.SEC.GUARATINGUETA", "DDM GUARATINGUETA"],
    "Lorena": ["", "01º D.P. LORENA", "02º D.P. LORENA", "DEL.POL.LORENA", "DDM LORENA"],
    "Piquete": ["", "DEL.POL.PIQUETE"],
    "Potim": ["", "DEL.POL.POTIM"],
    "Roseira": ["", "DEL.POL.ROSEIRA"]
}

# --- INTERFACE PRINCIPAL ---
st.title("Gerador de Laudos - Outras Peças")

st.header("1. Cabeçalho e Identificação")
colBO1, colBO2, colREP1, colREP2 = st.columns(4)
with colBO1: bo_input = st.text_input("Nº BO:", placeholder="Ex: DO4058-3", key=f"bo_{mk}").upper()
with colBO2: bo_ano = st.text_input("Ano BO:", value="2026", max_chars=4, key=f"ano_{mk}")
with colREP1: rep_input = st.text_input("Nº REP:", placeholder="Ex: 1234", key=f"rep_{mk}").upper()
with colREP2: rep_ano = st.text_input("Ano REP:", value="2026", max_chars=4, key=f"rep_ano_{mk}")

data_selecionada = st.date_input("Data do Laudo:", format="DD/MM/YYYY", key=f"data_{mk}")
meses = {1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril', 5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto', 9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'}
data_extenso = f"{data_selecionada.day} de {meses[data_selecionada.month]} de {data_selecionada.year}"

perito_selecionado = st.selectbox("Perito Criminal:", peritos, index=peritos.index("Claude Thiago Arrabal"), key=f"per_{mk}")
del_sel = st.selectbox("Autoridade Policial:", delegados, index=0, key=f"del_sel_{mk}")
delegado_selecionado = st.text_input("Digite o nome da Autoridade Policial:", key=f"del_dig_{mk}") if del_sel == "Outro..." else del_sel

colC1, colC2 = st.columns(2)
with colC1: 
    cid_sel = st.selectbox("Cidade:", cidades, index=0, key=f"cid_sel_{mk}")
    cidade_selecionada = st.text_input("Digite a Cidade:", key=f"cid_dig_{mk}") if cid_sel == "Outra..." else cid_sel
with colC2:
    if cid_sel == "Outra...": delegacia_selecionada = st.text_input("Digite a Delegacia:", key=f"dp_dig_{mk}")
    else:
        opcoes_dp = dps_por_cidade[cid_sel] + ["Outra..."]
        dp_sel = st.selectbox("Delegacia:", opcoes_dp, index=0, key=f"dp_sel_{mk}")
        delegacia_selecionada = st.text_input("Digite a Delegacia:", key=f"dp_dig_esp_{mk}") if dp_sel == "Outra..." else dp_sel

st.markdown("---")
st.header("2. Objetivo da Perícia")
obj_padrao = st.multiselect("Selecione os objetivos:", 
                            ["Fotografação", "Descrição", "Constatação de aptidão para corte de fios", "Constatação de aptidão para rompimento de obstáculo"], 
                            default=["Fotografação", "Descrição"])
obj_complemento = st.text_input("Complemento do Objetivo (Opcional):", placeholder="Ex: constatar características, estrutura e propriedades mecânicas...")

st.markdown("---")
st.header("3. Adicionar Itens Apreendidos")

with st.expander("➕ Clique aqui para adicionar um novo item", expanded=True):
    tipo_item = st.selectbox("Selecione o tipo de objeto:", ["Ferramenta Corta-Fio", "Marreta", "Pé de Cabra", "Outro Objeto"], key=f"tipo_item_{ik}")
    lacre_atual = st.text_input("Nº do Lacre de Entrada (Ex: 0072984):", key=f"lacre_ent_{ik}")
    
    # --- CORTA-FIO ---
    if tipo_item == "Ferramenta Corta-Fio":
        cor_cabo = st.text_input("Cor predominante dos cabos/empunhadura:", value="Amarela e preta", key=f"cor_{ik}")
        comprimento = st.number_input("Comprimento total aproximado (cm):", value=20.0, step=0.5, key=f"comp_{ik}")
        material = st.text_input("Material da área de corte:", value="Aço metálico", key=f"mat_{ik}")
        estado = st.selectbox("Estado de Conservação:", ["Bom", "Regular", "Com desgaste acentuado", "Péssimo"], key=f"estado_{ik}")
        teste = st.selectbox("Conclusão do Teste Prático:", ["Apta para o corte de fios metálicos", "Inapta para o corte de fios metálicos"], key=f"teste_{ik}")
        lacre_saida = st.text_input("Nº Lacre de Saída (Devolução):", key=f"lacre_sai_{ik}")

        if st.button("Adicionar Corta-Fio"):
            desc_fisica = f"Comprimento total aproximado de {comprimento} cm, com cabos revestidos na cor {cor_cabo} e área de corte em {material}. Estado de conservação {estado.lower()}."
            conclusao = f"A ferramenta foi submetida a testes práticos em laboratório, apresentando propriedades físicas e gumes cortantes que a tornam **{teste.lower()}** compatíveis com sua abertura."
            st.session_state['itens_pecas'].append({
                "lacre": lacre_atual, "lacre_saida": lacre_saida, "categoria": "Ferramenta Corta-Fio", 
                "descricao": desc_fisica, "conclusao": conclusao
            })
            st.success("Ferramenta adicionada!")

    # --- MARRETA ---
    elif tipo_item == "Marreta":
        peso = st.number_input("Peso aproximado (kg):", value=1.5, step=0.1, key=f"peso_{ik}")
        comp_cabo = st.number_input("Comprimento do cabo (cm):", value=30.0, step=1.0, key=f"comp_m_{ik}")
        mat_cabo = st.text_input("Material do cabo:", value="Madeira", key=f"mat_m_{ik}")
        estado = st.text_input("Estado/Desgaste:", value="Apresenta marcas de impacto e desgaste por uso", key=f"est_m_{ik}")
        teste = st.selectbox("Aptidão:", ["Apta para rompimento de obstáculos (percussão)", "Inapta para rompimento"], key=f"teste_m_{ik}")
        lacre_saida = st.text_input("Nº Lacre de Saída (Devolução):", key=f"lacre_sai_{ik}")

        if st.button("Adicionar Marreta"):
            desc_fisica = f"Constituída por cabeça de impacto de material metálico maciço, pesando aproximadamente {peso} kg, fixada a um cabo de {mat_cabo} medindo cerca de {comp_cabo} cm de comprimento. {estado}."
            conclusao = f"Reúne características de instrumento contundente, apresentando estrutura e resistência mecânica que a tornam **{teste.lower()}**."
            st.session_state['itens_pecas'].append({
                "lacre": lacre_atual, "lacre_saida": lacre_saida, "categoria": "Marreta", 
                "descricao": desc_fisica, "conclusao": conclusao
            })
            st.success("Marreta adicionada!")

    # --- PÉ DE CABRA ---
    elif tipo_item == "Pé de Cabra":
        comprimento = st.number_input("Comprimento aproximado (cm):", value=60.0, step=1.0, key=f"comp_p_{ik}")
        material = st.text_input("Material:", value="Vergalhão de aço metálico", key=f"mat_p_{ik}")
        extremidades = st.text_input("Formato das extremidades:", value="Uma extremidade achatada e outra curva", key=f"ext_{ik}")
        estado = st.text_input("Estado/Desgaste:", value="Apresenta oxidação e marcas de atrito", key=f"est_p_{ik}")
        teste = st.selectbox("Aptidão:", ["Apto para rompimento de obstáculos (alavanca)", "Inapto para rompimento"], key=f"teste_p_{ik}")
        lacre_saida = st.text_input("Nº Lacre de Saída (Devolução):", key=f"lacre_sai_{ik}")

        if st.button("Adicionar Pé de Cabra"):
            desc_fisica = f"Ferramenta tipo alavanca, confeccionada em {material}, medindo aproximadamente {comprimento} cm de comprimento. Apresenta {extremidades}. {estado}."
            conclusao = f"Reúne características de instrumento de alavancagem, apresentando estrutura e resistência mecânica que a tornam **{teste.lower()}**."
            st.session_state['itens_pecas'].append({
                "lacre": lacre_atual, "lacre_saida": lacre_saida, "categoria": "Pé de Cabra", 
                "descricao": desc_fisica, "conclusao": conclusao
            })
            st.success("Pé de cabra adicionado!")

    # --- OUTRO OBJETO ---
    elif tipo_item == "Outro Objeto":
        nome_obj = st.text_input("Nome do Objeto (Ex: Alicate de Pressão):", key=f"nome_out_{ik}")
        desc_livre = st.text_area("Descrição Física e Estado:", key=f"desc_out_{ik}")
        teste_livre = st.text_area("Conclusão de Testes/Aptidão (Opcional):", key=f"teste_out_{ik}")
        lacre_saida = st.text_input("Nº Lacre de Saída (Devolução):", key=f"lacre_sai_{ik}")

        if st.button("Adicionar Objeto"):
            st.session_state['itens_pecas'].append({
                "lacre": lacre_atual, "lacre_saida": lacre_saida, "categoria": nome_obj, 
                "descricao": desc_livre, "conclusao": teste_livre
            })
            st.success("Objeto adicionado!")

    # Botões de gerenciamento
    st.markdown("---")
    colB1, colB2 = st.columns(2)
    with colB1:
        if st.button("🆕 Novo Item (Limpar Formulário)", use_container_width=True):
            st.session_state['item_mk'] += 1
            st.rerun()
    with colB2:
        if st.button("🗑️ Remover Itens Adicionados", use_container_width=True):
            st.session_state['itens_pecas'] = []
            st.rerun()

# ==========================================
# 4. FOTOGRAFIAS 
# ==========================================
st.markdown("---")
st.header("4. Fotografias")

fotos_up = st.file_uploader("Arraste ou selecione as fotos aqui:", type=["jpg", "jpeg", "png"], accept_multiple_files=True, key=f"up_{st.session_state['uploader_key']}")

if fotos_up:
    nomes_atuais = [foto["nome"] for foto in st.session_state['fotos']]
    for f in fotos_up:
        if f.name not in nomes_atuais and f.name not in st.session_state['fotos_removidas']:
            img = Image.open(f)
            img = ImageOps.exif_transpose(img)
            st.session_state['fotos'].append({"imagem": img, "legenda": "Visão geral do objeto examinado.", "nome": f.name})
    
    st.session_state['uploader_key'] += 1
    st.rerun()

if st.session_state['fotos']:
    st.write("### Fotos Adicionadas:")
    for i, foto in enumerate(st.session_state['fotos']):
        cf1, cf2, cf3 = st.columns([1, 3, 1])
        with cf1: st.image(foto["imagem"], width=100)
        with cf2: 
            nova_leg = st.text_input("Legenda:", value=foto["legenda"], key=f"leg_{i}_{mk}")
            st.session_state['fotos'][i]["legenda"] = nova_leg
        with cf3:
            if st.button("Remover", key=f"rem_{i}_{mk}"):
                st.session_state['fotos_removidas'].append(st.session_state['fotos'][i]["nome"])
                st.session_state['fotos'].pop(i)
                st.rerun()

# ==========================================
# 5. EDIÇÃO FINAL DO TEXTO
# ==========================================
st.markdown("---")
st.header("5. Revisão e Edição de Texto")
st.info("💡 **Dica:** O texto gerado automaticamente aparece abaixo. Você pode reescrever, corrigir erros ou adicionar informações. Use **asteriscos duplos** para deixar uma palavra em negrito no Word (Ex: **Características:**).")

texto_obj_default = ", ".join(obj_padrao)
if obj_complemento: texto_obj_default += f", {obj_complemento}"
if texto_obj_default: texto_obj_default += "."

texto_exames_gerado = ""
if len(st.session_state['itens_pecas']) > 0:
    lacres_dict = {}
    for item in st.session_state['itens_pecas']:
        l = item.get('lacre', '')
        if l not in lacres_dict: lacres_dict[l] = []
        lacres_dict[l].append(item)

    contador_lacre = 1
    for lacre, lista_itens in lacres_dict.items():
        lacre_txt = f", com lacre de entrada nº {lacre}," if lacre else ","
        texto_exames_gerado += f"{contador_lacre}. Foi recebido para exames, em invólucro adequado{lacre_txt} os seguintes objetos:\n"
        for item in lista_itens:
            texto_exames_gerado += f"• **Tipo de Objeto:** {item.get('categoria', '')}.\n"
            texto_exames_gerado += f"• **Características Físicas:** {item.get('descricao', '')}\n"
            if item.get('conclusao'): 
                texto_exames_gerado += f"• **Testes e Conclusão:** {item.get('conclusao', '')}\n"
            if item.get('lacre_saida'): 
                texto_exames_gerado += f"O objeto foi acondicionado no lacre de saída nº {item.get('lacre_saida', '')}.\n"
            texto_exames_gerado += "\n"
        contador_lacre += 1

objetivo_final = st.text_area("Objetivo da Perícia (Editável):", value=texto_obj_default, height=68)
exames_final = st.text_area("Corpo dos Exames (Editável):", value=texto_exames_gerado, height=400)


# --- GERAÇÃO DO DOCUMENTO WORD ---
st.markdown("---")
if st.button("Criar Laudo (.docx)", type="primary", use_container_width=True):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Courier New'
    style.font.size = Pt(11)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    section = doc.sections[0]
    header = section.header
    for p in header.paragraphs: p.text = ""
    table = header.add_table(rows=1, cols=3, width=Cm(15.5))
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = table.columns[2].width = Cm(2.2); table.columns[1].width = Cm(11.1)
    for cell in table.columns[0].cells: cell.width = Cm(2.2)
    for cell in table.columns[1].cells: cell.width = Cm(11.1)
    for cell in table.columns[2].cells: cell.width = Cm(2.2)

    p_left = table.cell(0, 0).paragraphs[0]; p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists("logo_ssp.png"): p_left.add_run().add_picture("logo_ssp.png", width=Cm(1.8))
    
    p_center = table.cell(0, 1).paragraphs[0]; p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_h1 = p_center.add_run("SECRETARIA DA SEGURANÇA PÚBLICA\nSUPERINTENDÊNCIA DA POLÍCIA TÉCNICO-CIENTÍFICA\n")
    run_h1.font.size = Pt(11)
    run_h2 = p_center.add_run("INSTITUTO DE CRIMINALÍSTICA\n“PERITO CRIMINAL DR. OCTÁVIO EDUARDO DE BRITO ALVARENGA”\nNÚCLEO DE PERÍCIAS CRIMINALÍSTICAS DE SÃO JOSÉ DOS CAMPOS\nEQUIPE DE PERÍCIAS CRIMINALÍSTICAS DE GUARATINGUETÁ")
    run_h2.font.size = Pt(8)

    p_right = table.cell(0, 2).paragraphs[0]; p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists("logo_ic.png"): p_right.add_run().add_picture("logo_ic.png", width=Cm(1.8))

    if rep_input or bo_input:
        p_id = doc.add_paragraph()
        texto_id = []
        if rep_input: texto_id.append(f"REP {rep_input.upper()} / {rep_ano}")
        if bo_input: texto_id.append(f"BO {bo_input.upper()} / {bo_ano}")
        
        delegacia_txt = f" - {delegacia_selecionada}" if delegacia_selecionada else ""
        p_id.add_run(" - ".join(texto_id) + delegacia_txt)
        p_id.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # CORPO
    p_nat = doc.add_paragraph()
    # 1. NATUREZA ALTERADA PARA PEÇAS
    run = p_nat.add_run("1 – NATUREZA: Exame em Peças (Objetos e Ferramentas)"); run.bold = True; run.font.size = Pt(14)
    adicionar_borda_inferior(p_nat)
    
    del_nome = delegado_selecionado if delegado_selecionado else "__________"
    del_dp = delegacia_selecionada if delegacia_selecionada else "__________"
    preambulo = (f"Aos {data_extenso}, no Instituto de Criminalística, da Superintendência da Polícia Técnico-Científica, "
                 f"da Secretaria da Segurança Pública do Estado de São Paulo, de conformidade com o disposto no artigo 178 "
                 f"do Decreto-Lei nº. 3689, de 03 de outubro de 1941, pelo Diretor do Instituto de Criminalística, Ricardo Lopes Ortega, "
                 f"foi designado o Perito Criminal {perito_selecionado}, para proceder ao exame supracitado, em atendimento à requisição "
                 f"da Autoridade Policial, Dr(a). {del_nome}, titular/em exercício na {del_dp}.")
    doc.add_paragraph(preambulo)

    p_obj = doc.add_paragraph()
    run_obj = p_obj.add_run("2 - OBJETIVO DA PERÍCIA:"); run_obj.bold = True; run_obj.font.size = Pt(14)
    adicionar_borda_inferior(p_obj)
    doc.add_paragraph(objetivo_final)

    p_ex = doc.add_paragraph()
    run_ex = p_ex.add_run("3 – DOS EXAMES:"); run_ex.bold = True; run_ex.font.size = Pt(14)
    adicionar_borda_inferior(p_ex)

    for linha in exames_final.split('\n'):
        if not linha.strip():
            doc.add_paragraph()
            continue
            
        if linha.startswith("• "):
            p = doc.add_paragraph(style='List Bullet')
            linha_limpa = linha[2:]
        else:
            p = doc.add_paragraph()
            linha_limpa = linha
            
        partes = re.split(r'(\*\*.*?\*\*)', linha_limpa)
        for parte in partes:
            if parte.startswith('**') and parte.endswith('**'):
                p.add_run(parte[2:-2]).bold = True
            else:
                p.add_run(parte)

    # ILUSTRATIVO FOTOGRÁFICO
    if st.session_state['fotos']:
        doc.add_page_break()
        p_foto_cabecalho = doc.add_paragraph()
        run_foto = p_foto_cabecalho.add_run("4 – DO ILUSTRATIVO FOTOGRÁFICO:"); run_foto.bold = True; run_foto.font.size = Pt(14)
        adicionar_borda_inferior(p_foto_cabecalho)
        
        for foto in st.session_state['fotos']:
            img_stream = io.BytesIO()
            foto['imagem'].save(img_stream, format='PNG')
            img_stream.seek(0)
            
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = p_img.add_run()
            run_img.add_picture(img_stream, width=Cm(14))
            
            p_leg = doc.add_paragraph(f"Figura: {foto['legenda']}")
            p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ENCERRAMENTO PADRÃO
    p_relatar = doc.add_paragraph("Era o que havia a relatar.")
    p_relatar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_paginas = doc.add_paragraph("Este laudo vai impresso em ")
    adicionar_campo_numpages(p_paginas)  
    p_paginas.add_run(" páginas, além da capa, ficando arquivada cópia digital no sistema GDL da SPTC.")
    p_paginas.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_assinatura = doc.add_paragraph(); p_assinatura.paragraph_format.space_after = p_assinatura.paragraph_format.space_before = Pt(0)
    p_assinatura.add_run(perito_selecionado).bold = True
    p_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_cargo = doc.add_paragraph("Perito Criminal Relator"); p_cargo.paragraph_format.space_after = p_cargo.paragraph_format.space_before = Pt(0)
    p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf_doc = io.BytesIO(); doc.save(buf_doc); buf_doc.seek(0)
    
    if rep_input:
        nome_arquivo = f"Laudo_Pecas_REP_{rep_input}_{rep_ano}.docx"
    elif bo_input:
        nome_arquivo = f"Laudo_Pecas_BO_{bo_input}_{bo_ano}.docx"
    else:
        nome_arquivo = "Laudo_Pecas_Sem_BO_REP.docx"
        
    st.download_button("⬇️ Descarregar Laudo Final", buf_doc, nome_arquivo, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)