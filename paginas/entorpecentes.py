import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import io
import os
import json
from datetime import datetime, timedelta

# --- Funções de Formatação do Word ---
def adicionar_campo_numpages(paragraph):
    p = paragraph._p
    r1 = OxmlElement('w:r')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    r1.append(fldChar1)
    p.append(r1)
    r2 = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' NUMPAGES '
    r2.append(instrText)
    p.append(r2)
    r3 = OxmlElement('w:r')
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    r3.append(fldChar2)
    p.append(r3)
    r4 = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = 'xx'
    r4.append(t)
    p.append(r4)
    r5 = OxmlElement('w:r')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r5.append(fldChar3)
    p.append(r5)

def aplicar_marca_texto(run, cor):
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), cor)
    rPr.append(highlight)

# --- Variáveis de Sessão ---
if 'mk' not in st.session_state: st.session_state['mk'] = 0 
mk = st.session_state['mk']

# --- Dicionários e Listas do VBA ---
te_map = {
    "Invólucro Geral": "Invólucro(s) plástico(s) encerrando",
    "Eppendorf": 'Microtubo(s) plástico(s) do tipo "Eppendorf" dotado(s) de tampa própria encerrando',
    "Plástico Nó": "Invólucro(s) plástico(s) fechado(s) por nó encerrando",
    "Plástico Filme": 'Invólucro(s) plástico(s) do tipo "filme" retorcido encerrando',
    "Plástico Zip": 'Invólucro(s) plástico(s) fechado(s) por pressão (tipo "zip") encerrando',
    "Plástico Calor": "Invólucro(s) plástico(s) fechado(s) por aquecimento encerrando",
    "Plástico Alumínio": "Invólucro(s) constituído(s) por plástico e papel alumínio encerrando",
    "Fita Adesiva": "Invólucro(s) constituído(s) por fita(s) adesiva(s) retorcida(s) encerrando",
    "Frasco Vítreo": "Frasco(s) vítreos dotado(s) de válvula aspersora contendo"
}

sb_map = {
    "Maconha": "porção de fragmentos vegetais ressequidos, constituídos de folhas, folíolos, inflorescências, caules e frutos.",
    "Cocaína": "porção de material sólido particulado.",
    "Crack": "porção de material sólido petrificado.",
    "Tijolo": "porção de fragmentos vegetais ressequidos, constituídos de folhas, folíolos, inflorescências, caules e frutos, compactados na forma de tijolo.",
    "Haxixe": "porção de substância de aspecto resinoso de coloração amarronzada e de morfologia irregular.",
    "Resina": "porção de material resinoso.",
    "Cigarro Íntegro": "cigarro(s) artesanal(ais) confeccionado(s) em papel, contendo fragmentos vegetais ressequidos, constituídos de folhas, folíolos, inflorescências, caules e frutos.",
    "Cigarro Queimado": "cigarro(s) artesanal(ais) parcialmente queimado(s), confeccionado(s) em papel, contendo fragmentos vegetais ressequidos, constituídos de folhas, folíolos, inflorescências, caules e frutos."
}

delegados = ["Selecione..."] + sorted(["Adilson Antonio Marcondes dos Santos", "Adriane Goncalves", "Anisio Galdioli", "Benedito Carlos dos Santos Martins", "Cesar Aparecido Vieira da Silva", "Cristiane Correa de Freitas", "Daniel Souza Baptista de Castro", "Ernani Ronaldo Giannico Braga", "Fabio Germano Figueiredo Cabett", "Flavia Maria Rocha Rollo", "Francisco Sannini Neto", "Hugo Parreiras de Macedo", "Jose Marcelo Silva Hial", "Leonardo da Costa Ferreira", "Marcelo Vieira Cavalcante", "Mario Celso Ribeiro Senne", "Paulo Roberto Gruschka Castilho", "Paulo Sergio Barbosa", "Pedro Rossati", "Rodrigo Jose Goes Ribeiro", "Sergio Lucas Adler Guedes de Oliveira", "Vania Idalira Z. de Oliveira"]) + ["Outro..."]
peritos = ["Alexandre Rabello de Oliveira", "Bruna Fernandes Nogueira", "Claude Thiago Arrabal", "Jéssica Pereira Gonçalves", "Júlia Soares Melo", "Luiz Fausto Prado Vasques", "Luiza Dias da Cunha Lima", "Marcelo Mourão Dantas", "Márcio Steinmetz Soares", "Rafael Rodrigues Cunha", "Sarah Costa Teixeira", "Ruan Carvalho de Souza"]
cidades = ["Selecione...", "Aparecida", "Cachoeira Paulista", "Canas", "Cunha", "Guaratinguetá", "Lorena", "Piquete", "Potim", "Roseira", "Outra..."]

dps_por_cidade = {
    "Aparecida": ["DEL.POL.APARECIDA"],
    "Canas": ["DEL.POL.CANAS"],
    "Cachoeira Paulista": ["DEL.POL.CACHOEIRA PAULISTA"],
    "Cunha": ["DEL.POL.CUNHA"],
    "Guaratinguetá": ["01º D.P. GUARATINGUETA", "02º D.P. GUARATINGUETA", "03º D.P. GUARATINGUETA", "DEL.INV.GER. GUARATINGUETA", "DEL.SEC.GUARATINGUETA", "DEL.SEC.GUARATINGUETA PLANTÃO", "DISE- DEL.SEC.GUARATINGUETA"],
    "Lorena": ["01º D.P. LORENA", "02º D.P. LORENA", "DEL.POL.LORENA"],
    "Piquete": ["DEL.POL.PIQUETE"],
    "Potim": ["DEL.POL.POTIM"],
    "Roseira": ["DEL.POL.ROSEIRA"]
}

# --- INTERFACE ---
st.title("Gerador de Laudos - Entorpecentes")

st.header("1. Cabeçalho e Identificação")

col1, col2 = st.columns(2)
with col1:
    bo_input = st.text_input("Número do BO:", key=f"bo_{mk}").upper()
    rep_input = st.text_input("Número REP (Se houver):", key=f"rep_{mk}")
    investigado = st.text_input("Investigado (Opcional):", key=f"inv_{mk}")
    
with col2:
    bo_ano = st.text_input("Ano do BO:", value="2026", key=f"ano_{mk}")
    rep_ano = st.text_input("Ano REP:", value="2026", key=f"rep_ano_{mk}")
    data_selecionada = st.date_input("Data do Laudo:", format="DD/MM/YYYY", key=f"data_{mk}")

meses = {1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril', 5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto', 9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'}
data_extenso = f"{data_selecionada.day} de {meses[data_selecionada.month]} de {data_selecionada.year}"

# Local do Fato
st.markdown("**Local do Fato**")
cF1, cF2, cF3 = st.columns([2, 1, 1])
with cF1: endereco = st.text_input("Endereço:", key=f"end_{mk}")
with cF2: numero = st.text_input("Número:", key=f"num_{mk}")
with cF3: complemento = st.text_input("Bairro/Comp.:", key=f"comp_{mk}")

perito_selecionado = st.selectbox("Perito Criminal:", peritos, index=peritos.index("Claude Thiago Arrabal"), key=f"per_{mk}")

# Credenciais GDL para Exportação
with st.expander("🔑 Credenciais para Automação GDL"):
    cCred1, cCred2 = st.columns(2)
    with cCred1: user_gdl = st.text_input("Usuário GDL:", value="claude.cta", key=f"usr_gdl_{mk}")
    with cCred2: pass_gdl = st.text_input("Senha GDL:", type="password", value="WQA4M1", key=f"pwd_gdl_{mk}")

# Autoridade
del_sel = st.selectbox("Autoridade Policial:", delegados, index=0, key=f"del_sel_{mk}")
delegado_selecionado = st.text_input("Digite o nome da Autoridade:", key=f"del_dig_{mk}") if del_sel == "Outro..." else del_sel

# Cidade e DP
colC1, colC2 = st.columns(2)
with colC1: 
    cid_sel = st.selectbox("Cidade:", cidades, index=0, key=f"cid_sel_{mk}")
with colC2:
    if cid_sel == "Selecione...":
        dp_sel = st.selectbox("Delegacia:", ["Selecione a cidade primeiro..."], disabled=True, key=f"dp_sel_disabled_{mk}")
        delegacia_selecionada = "[DELEGACIA NÃO INFORMADA]"
    elif cid_sel in dps_por_cidade:
        opcoes_dp = ["Selecione..."] + dps_por_cidade[cid_sel] + ["Outra..."]
        dp_sel = st.selectbox("Delegacia:", opcoes_dp, index=0, key=f"dp_sel_{mk}")
    else:
        dp_sel = st.text_input("Digite a Delegacia:", key=f"dp_dig_{mk}")

if cid_sel != "Selecione...":
    delegacia_selecionada = st.text_input("Nome da DP Específica:", key=f"dp_esp_{mk}") if dp_sel == "Outra..." else dp_sel

st.header("2. Materiais e Lacres (Por Item)")

colL1, colL2 = st.columns(2)
with colL1:
    lacre_saida_delegacia = st.text_input("Lacre de Saída Delegacia:", help="Será copiado para o Lacre Saída (Devolvido) de todos os itens.", key=f"ls_geral_{mk}")
with colL2:
    qtd_itens = st.number_input("Quantidade de Itens:", min_value=1, max_value=20, value=1, step=1, key=f"qtd_{mk}")

dados_itens = []
massa_total = 0.0

for i in range(qtd_itens):
    st.markdown(f"### 🌿 Item {i+1}")
    
    ci1, ci2, ci3 = st.columns(3)
    with ci1: emb = st.selectbox(f"Embalagem (Item {i+1}):", list(te_map.keys()), key=f"emb_{i}_{mk}")
    with ci2: sub = st.selectbox(f"Substância (Item {i+1}):", list(sb_map.keys()), key=f"sub_{i}_{mk}")
    with ci3: res = st.selectbox(f"Resultado (Item {i+1}):", ["Positivo", "Negativo", "Inconclusivo"], key=f"res_{i}_{mk}")
    
    cm1, cm2, cm3 = st.columns(3)
    # Step=None remove os botões de + e -
    with cm1: mb = st.number_input(f"Massa Bruta (g)", min_value=0.00, value=1.00, format="%.2f", step=None, key=f"mb_{i}_{mk}")
    with cm2: ml = st.number_input(f"Massa Líquida (g)", min_value=0.00, value=1.00, format="%.2f", step=None, key=f"ml_{i}_{mk}")
    with cm3: am = st.number_input(f"Amostra (g)", min_value=0.00, value=2.00, format="%.2f", step=None, key=f"am_{i}_{mk}")
    
    cl1, cl2, cl3 = st.columns(3)
    with cl1: le = st.text_input(f"Lacre Entrada", key=f"le_{i}_{mk}")
    with cl2: ld = st.text_input(f"Lacre Saída (Devolvido)", value=lacre_saida_delegacia, key=f"ld_{i}_{mk}")
    with cl3: ls = st.text_input(f"Lacre Saída (CP - Amostra)", key=f"ls_{i}_{mk}")

    massa_total += mb

    dados_itens.append({
        "emb": emb, "sub": sub, "res": res, 
        "mb": mb, "ml": ml, "am": am, 
        "le": le, "ld": ld, "ls": ls
    })


# --- GERAÇÃO DE JSON PARA GDL ---
hora_lib = (datetime.now() - timedelta(minutes=10)).strftime("%H%M")
cidade_export = "Guaratinguet" if cid_sel == "Guaratinguetá" else (cid_sel if cid_sel != "Selecione..." else "")

dp_export = delegacia_selecionada
if dp_export.startswith("01") or dp_export.startswith("02") or dp_export.startswith("03"):
    dp_export = dp_export[:2]

dados_exportacao = {
    "usuario": user_gdl,
    "senha": pass_gdl,
    "delegado": delegado_selecionado if delegado_selecionado != "[AUTORIDADE NÃO INFORMADA]" else "",
    "bo": bo_input,
    "cidade": cidade_export,
    "dp": dp_export,
    "hora_liberacao": hora_lib,
    "cidade_fato": cidade_export,
    "investigado": investigado,
    "endereco": endereco,
    "numero": numero,
    "complemento": complemento,
    "itens": []
}

for ap in dados_itens:
    dados_exportacao["itens"].append({
        "substancia": ap['sub'], 
        "massa": f"{ap['ml']:.2f}".replace(".", ","), # GDL usa vírgula
        "lacre_entrada": ap['le'],
        "lacre_saida": ap['ld']
    })

json_texto = json.dumps(dados_exportacao, ensure_ascii=False, indent=4)

# --- FINALIZAÇÃO E DOWNLOAD ---
st.header("3. Finalizar e Exportar")
c1, c2, c3 = st.columns(3)

with c1:
    if st.button("📄 GERAR LAUDO (.DOCX)", type="primary", use_container_width=True):
        doc = Document()
        
        # Formatando o estilo geral para não ter espaços extras
        style = doc.styles['Normal']
        style.font.name = 'Courier New'
        style.font.size = Pt(10)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        
        # --- CABEÇALHO SPTC ---
        section = doc.sections[0]
        header = section.header
        for p in header.paragraphs: p.text = ""
        table = header.add_table(rows=1, cols=3, width=Cm(15.5))
        table.columns[0].width = table.columns[2].width = Cm(2.2); table.columns[1].width = Cm(11.1)
        
        # Limpando espaços dos parágrafos da tabela
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)

        p_left = table.cell(0, 0).paragraphs[0]; p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if os.path.exists("logo_ssp.png"): p_left.add_run().add_picture("logo_ssp.png", width=Cm(1.8))
        
        p_c = table.cell(0, 1).paragraphs[0]; p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_h1 = p_c.add_run("SECRETARIA DA SEGURANÇA PÚBLICA\nSUPERINTENDÊNCIA DA POLÍCIA TÉCNICO-CIENTÍFICA\n")
        run_h1.font.size = Pt(11)
        run_h2 = p_c.add_run("INSTITUTO DE CRIMINALÍSTICA\nNÚCLEO DE PERÍCIAS CRIMINALÍSTICAS DE SÃO JOSÉ DOS CAMPOS\nEQUIPE DE PERÍCIAS CRIMINALÍSTICAS DE GUARATINGUETÁ")
        run_h2.font.size = Pt(8)
        
        p_right = table.cell(0, 2).paragraphs[0]; p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if os.path.exists("logo_ic.png"): p_right.add_run().add_picture("logo_ic.png", width=Cm(1.8))

        # --- IDENTIFICAÇÃO (BO / REP) ---
        if rep_input or bo_input:
            p_ident = doc.add_paragraph()
            p_ident.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ident_text = []
            if rep_input: ident_text.append(f"REP {rep_input} / {rep_ano}")
            if bo_input: ident_text.append(f"BO {bo_input} / {bo_ano} - {delegacia_selecionada}")
            p_ident.add_run(" | ".join(ident_text))

        # --- TEXTOS DO LAUDO ---
        p_nat = doc.add_paragraph()
        run = p_nat.add_run("Natureza do Exame: "); run.bold = True
        p_nat.add_run("Constatação Provisória de Entorpecentes.")
        
        preambulo = (f"Aos {data_extenso}, no Instituto de Criminalística, da Superintendência da Técnico-Científica, "
                     f"da Secretaria da Segurança Pública do Estado de São Paulo, de conformidade Polícia com o disposto "
                     f"no artigo 178 do Decreto-Lei nº. 3689, de 03 de outubro de 1941, pelo Diretor do Instituto de Criminalística, "
                     f"Dr. Ricardo Lopes Ortega, foi designado o Perito Criminal {perito_selecionado}, para proceder ao exame supracitado, "
                     f"em atendimento à requisição subscrita pelo Dr(a). {delegado_selecionado}, Delegado(a) de Polícia.")
        doc.add_paragraph(preambulo)
        doc.add_paragraph("") # Pula linha
        
        p_mat = doc.add_paragraph()
        p_mat.add_run(f"Dos Materiais Recebidos e Examinados ({qtd_itens} Item(s)):").bold = True
        doc.add_paragraph("Todo material recebido encontrava-se acondicionado em invólucro(s) plástico(s) lacrado(s), acompanhado da requisição de exame pericial.")
        
        p_mt = doc.add_paragraph()
        p_mt.add_run("Massa Bruta Apresentada: ").bold = True
        p_mt.add_run(f"{massa_total:.2f} grama(s).".replace(".", ",")) # Ponto no final
        doc.add_paragraph("") # Pula linha

        # Loop de Itens
        for i, item in enumerate(dados_itens):
            p_item = doc.add_paragraph()
            p_item.add_run(f"Item {i+1} (Acondicionado sob o lacre {item['le']}).").bold = True
            
            p_desc = doc.add_paragraph()
            p_desc.add_run("Descrição: ").bold = True
            p_desc.add_run(f"{te_map[item['emb']]} {sb_map[item['sub']]}")
            
            p_massa = doc.add_paragraph()
            p_massa.add_run("Massa Bruta e/ou Quantidade: ").bold = True
            p_massa.add_run(f"{item['mb']:.2f} grama(s). ".replace(".", ",")) # Ponto e espaço
            p_massa.add_run("Massa Líquida: ").bold = True
            p_massa.add_run(f"{item['ml']:.2f} grama(s).".replace(".", ",")) # Ponto no final
            doc.add_paragraph("") # Pula linha

            # Amostra
            doc.add_paragraph(f"Uma amostra de {item['am']:.2f} grama(s) foi aqui retirada para análises. O material remanescente destas análises foi fechado sob o lacre de número {item['ls']} e será encaminhado ao IC - CP - São Jose dos Campos para a elaboração do respectivo Laudo Definitivo, complementar a este Laudo de Constatação.".replace(".", ","))
            doc.add_paragraph("") # Pula linha

            # Resultado
            p_res = doc.add_paragraph()
            if item['res'] == "Positivo":
                if item['sub'] in ["Crack", "Cocaína"]:
                    p_res.add_run("A análise do material descrito fez o uso de teste colorimétrico sendo ")
                    r_hl = p_res.add_run("DETECTADA presença da substância COCAÍNA")
                    aplicar_marca_texto(r_hl, "green")
                    p_res.add_run(", constante na lista F1 da Portaria SVS/MS 344/98 e atualizações posteriores.")
                else: 
                    p_res.add_run("A análise do material descrito fez o uso de teste colorimétrico sendo ")
                    r_hl = p_res.add_run("DETECTADA presença da substância TETRAHIDROCANNABINOL (THC)")
                    aplicar_marca_texto(r_hl, "green")
                    p_res.add_run(", constante na lista F2 da Portaria SVS/MS 344/98 e atualizações posteriores.")
            
            elif item['res'] == "Negativo":
                p_res.add_run("A análise do material descrito fez o uso de teste colorimétrico e ")
                r_hl = p_res.add_run("NÃO FOI POSSÍVEL IDENTIFICAR")
                aplicar_marca_texto(r_hl, "red")
                p_res.add_run(" presença de substâncias elencadas nas listas A, B e F da Portaria SVS/MS 344/98 e atualizações posteriores, ou na Portaria MJSP 204/2022, em sua lista III, conforme a(s) técnica(s) utilizada(s) (Portaria SPTC 42/2024).")
            
            elif item['res'] == "Inconclusivo":
                p_res.add_run("A análise do material descrito fez o uso de teste colorimétrico. Os exames/análises preliminares mostraram-se ")
                r_hl = p_res.add_run("INCONCLUSIVOS")
                aplicar_marca_texto(r_hl, "yellow")
                p_res.add_run(", sendo necessárias análises mais complexas e morosas, incompatíveis com a rapidez demandada pelos exames de constatação. O resultado deste presente item seguirá em laudo definitivo.")
            
            doc.add_paragraph("") # Pula linha
            doc.add_paragraph(f"O restante do item (material, invólucro(s) e lacre(s)) foi devolvido à autoridade policial requisitante nos termos das exigências legais, sob o lacre número {item['ld']}.")
            doc.add_paragraph("") # Pula linha

        # Fechamento e Assinatura
        p_assinatura = doc.add_paragraph()
        p_assinatura.add_run(perito_selecionado)
        p_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_cargo = doc.add_paragraph()
        p_cargo.add_run("Perito Criminal Relator").bold = True
        p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        
        # Nome do arquivo prioriza BO se não tiver REP
        if bo_input:
            nome_arquivo = f"Laudo_Ent_BO_{bo_input}_{bo_ano}.docx"
        elif rep_input:
            nome_arquivo = f"Laudo_Ent_REP_{rep_input}_{rep_ano}.docx"
        else:
            nome_arquivo = "Laudo_Entorpecentes.docx"
            
        st.download_button("⬇️ Salvar Laudo", buf, nome_arquivo, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

with c2:
    st.download_button(
        label="🚀 Baixar Dados GDL (JSON)",
        data=json_texto,
        file_name="dados_gdl.json",
        mime="application/json",
        use_container_width=True
    )

with c3:
    if st.button("🔄 Novo Laudo (Limpar)", type="secondary", use_container_width=True):
        st.session_state.clear()
        st.session_state['mk'] = mk + 1
        st.rerun()