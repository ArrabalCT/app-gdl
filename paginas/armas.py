import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import io
import os
import re
from PIL import Image, ImageOps

# --- Funções de Formatação do Word ---
def adicionar_borda_inferior(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def adicionar_campo_numpages(paragraph):
    p = paragraph._p
    for char_type, text in [('begin', None), (None, ' NUMPAGES '), ('separate', None), (None, 'xx'), ('end', None)]:
        r = OxmlElement('w:r')
        if char_type:
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), char_type)
            r.append(fldChar)
        else:
            instrText = OxmlElement('w:instrText' if 'NUMPAGES' in text else 'w:t')
            if 'NUMPAGES' in text: instrText.set(qn('xml:space'), 'preserve')
            instrText.text = text
            r.append(instrText)
        p.append(r)

# --- BASE DE DADOS DE PROJÉTEIS (CBC) ---
TABELA_PROJETEIS = [
    {"calibre": ".25 Auto", "diametro": 6.3, "massa": 3.2},
    {"calibre": ".32 Auto", "diametro": 7.8, "massa": 4.6},
    {"calibre": ".32 S&W", "diametro": 7.9, "massa": 6.3},
    {"calibre": ".380 Auto", "diametro": 9.0, "massa": 6.1},
    {"calibre": "9mm Luger", "diametro": 9.0, "massa": 7.4},
    {"calibre": "9mm Luger", "diametro": 9.0, "massa": 8.0},
    {"calibre": ".38 SPL", "diametro": 9.0, "massa": 10.2},
    {"calibre": ".357 Magnum", "diametro": 9.0, "massa": 10.2},
    {"calibre": ".40 S&W", "diametro": 10.1, "massa": 10.0},
    {"calibre": ".44-40", "diametro": 10.8, "massa": 12.9},
    {"calibre": ".45 Auto", "diametro": 11.4, "massa": 14.9},
    {"calibre": ".44 Magnum", "diametro": 11.4, "massa": 15.5},
    {"calibre": ".308", "diametro": 7.8, "massa": 9.7},
    {"calibre": ".223", "diametro": 5.6, "massa": 3.6},
]

def estimar_calibre(massa, diametro):
    resultados = []
    for p in TABELA_PROJETEIS:
        if abs(p["diametro"] - diametro) <= 0.2 and abs(p["massa"] - massa) <= 0.5:
            if p["calibre"] not in resultados: resultados.append(p["calibre"])
    return " ou ".join(resultados) if resultados else "Não identificado na tabela padrão"

# --- Variáveis de Sessão ---
if 'itens_balistica' not in st.session_state: st.session_state['itens_balistica'] = []
if 'fotos' not in st.session_state: st.session_state['fotos'] = []
if 'fotos_removidas' not in st.session_state: st.session_state['fotos_removidas'] = []
if 'mk' not in st.session_state: st.session_state['mk'] = 0 
if 'uploader_key' not in st.session_state: st.session_state['uploader_key'] = 0
if 'item_mk' not in st.session_state: st.session_state['item_mk'] = 0 # Variável para o botão Novo Item
mk = st.session_state['mk']
ik = st.session_state['item_mk']

# --- Listas de Dados ---
delegados = ["", "Adilson Antonio Marcondes dos Santos", "Adriane Goncalves", "Anisio Galdioli", "Cesar Aparecido Vieira da Silva", "Daniel Souza Baptista de Castro", "Ernani Ronaldo Giannico Braga", "Fabio Germano Figueiredo Cabett", "Flavia Maria Rocha Rollo", "Francisco Sannini Neto", "Hugo Parreiras de Macedo", "Jose Marcelo Silva Hial", "Leonardo da Costa Ferreira", "Marcelo Vieira Cavalcante", "Mario Celso Ribeiro Senne", "Paulo Roberto Gruschka Castilho", "Paulo Sergio Barbosa", "Pedro Rossati", "Sergio Lucas Adler Guedes de Oliveira", "Vania Idalira Z. de Oliveira", "Outro..."]
peritos = ["Alexandre Rabello de Oliveira", "Bruna Fernandes Nogueira", "Claude Thiago Arrabal", "Jéssica Pereira Gonçalves", "Júlia Soares Melo", "Luiz Fausto Prado Vasques", "Marcelo Mourão Dantas", "Márcio Steinmetz Soares", "Sarah Costa Teixeira", "Ruan Carvalho de Souza"]
cidades = ["", "Aparecida", "Cachoeira Paulista", "Canas", "Cunha", "Guaratinguetá", "Lorena", "Piquete", "Potim", "Roseira", "Outra..."]
dps_por_cidade = {
    "": [""],
    "Aparecida": ["", "DEL.POL.APARECIDA", "01º D.P. APARECIDA"],
    "Canas": ["", "DEL.POL.CANAS"],
    "Cachoeira Paulista": ["", "DEL.POL.CACHOEIRA PAULISTA", "01º D.P. CACHOEIRA PAULISTA"],
    "Cunha": ["", "DEL.POL.CUNHA"],
    "Guaratinguetá": ["", "01º D.P. GUARATINGUETA", "02º D.P. GUARATINGUETA", "03º D.P. GUARATINGUETA", "DEL.SEC.GUARATINGUETA PLANTÃO", "DISE- DEL.SEC.GUARATINGUETA", "DDM GUARATINGUETA"],
    "Lorena": ["", "01º D.P. LORENA", "02º D.P. LORENA", "DEL.POL.LORENA", "DDM LORENA"],
    "Piquete": ["", "DEL.POL.PIQUETE"],
    "Potim": ["", "DEL.POL.POTIM"],
    "Roseira": ["", "DEL.POL.ROSEIRA"]
}

# --- INTERFACE PRINCIPAL ---
st.title("Gerador de Laudos - Armas e Munições")

st.header("1. Cabeçalho e Identificação")
colBO1, colBO2, colREP1, colREP2 = st.columns(4)
with colBO1: bo_input = st.text_input("Nº BO:", placeholder="Ex: LT0644", key=f"bo_{mk}").upper()
with colBO2: bo_ano = st.text_input("Ano BO:", value="2026", max_chars=4, key=f"ano_{mk}")
with colREP1: rep_input = st.text_input("Nº REP:", placeholder="Ex: 1234", key=f"rep_{mk}").upper()
with colREP2: rep_ano = st.text_input("Ano REP:", value="2026", max_chars=4, key=f"rep_ano_{mk}")

data_selecionada = st.date_input("Data do Laudo:", format="DD/MM/YYYY", key=f"data_{mk}")
meses = {1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril', 5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto', 9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'}
data_extenso = f"{data_selecionada.day} de {meses[data_selecionada.month]} de {data_selecionada.year}"

perito_selecionado = st.selectbox("Perito Criminal:", peritos, index=peritos.index("Claude Thiago Arrabal"), key=f"per_{mk}")
del_sel = st.selectbox("Autoridade Policial:", delegados, index=0, key=f"del_sel_{mk}")
delegado_selecionado = st.text_input("Digite o nome da Autoridade Policial:", key=f"del_dig_{mk}") if del_sel == "Outro..." else del_sel

colC1, colC2 = st.columns(2)
with colC1: 
    cid_sel = st.selectbox("Cidade:", cidades, index=0, key=f"cid_sel_{mk}")
    cidade_selecionada = st.text_input("Digite a Cidade:", key=f"cid_dig_{mk}") if cid_sel == "Outra..." else cid_sel
with colC2:
    if cid_sel == "Outra...": delegacia_selecionada = st.text_input("Digite a Delegacia:", key=f"dp_dig_{mk}")
    else:
        opcoes_dp = dps_por_cidade[cid_sel] + ["Outra..."]
        dp_sel = st.selectbox("Delegacia:", opcoes_dp, index=0, key=f"dp_sel_{mk}")
        delegacia_selecionada = st.text_input("Digite a Delegacia:", key=f"dp_dig_esp_{mk}") if dp_sel == "Outra..." else dp_sel

st.markdown("---")
st.header("2. Objetivo da Perícia")
obj_padrao = st.multiselect("Selecione os objetivos:", 
                            ["Fotografação", "Descrição", "Calibre", "Eficácia", "Recenticidade", "Potencialidade Lesiva"], 
                            default=["Fotografação", "Descrição", "Calibre", "Eficácia"])
obj_complemento = st.text_input("Complemento do Objetivo (Opcional):", placeholder="Ex: constatação de numeração suprimida...")


st.markdown("---")
st.header("3. Adicionar Itens Apreendidos")

with st.expander("➕ Clique aqui para adicionar um novo item", expanded=True):
    tipo_item = st.selectbox("O que você vai adicionar?", ["Arma de Fogo", "Munições", "Estojos", "Projétil"], key=f"tipo_item_{ik}")
    lacre_atual = st.text_input("Nº do Lacre de Entrada (Ex: 00004041):", key=f"lacre_ent_{ik}")
    
    # --- ARMA ---
    if tipo_item == "Arma de Fogo":
        t_sel = st.selectbox("Tipo da Arma:", ["", "PISTOLA", "REVÓLVER", "ESPINGARDA", "CARABINA", "FUZIL", "GARRUCHA", "ARTESANAL", "Outra..."], key=f"t_sel_{ik}")
        tipo_arma = st.text_input("Especifique o tipo:", key=f"t_esp_{ik}") if t_sel == "Outra..." else t_sel
        fab_arma = st.text_input("Fabricante / Modelo:", key=f"fab_arma_{ik}")
        cal_arma = st.text_input("Calibre Nominal:", key=f"cal_arma_{ik}")
        estado_arma = st.selectbox("Estado de Conservação:", ["Bom", "Regular", "Ruim"], key=f"estado_arma_{ik}")
        
        st.write("**Características Físicas e Componentes:**")
        
        detalhes_arma = ""
        if t_sel in ["PISTOLA", "FUZIL", "CARABINA", "ARTESANAL"]:
            tem_carr = st.selectbox("Acompanha carregador?", ["", "Sim", "Não"], key=f"tem_carr_{ik}")
            if tem_carr == "Sim":
                qtd_carr = st.number_input("Quantos carregadores?", min_value=1, value=1, key=f"qtd_carr_{ik}")
                cap_carr = st.number_input("Capacidade do carregador (munições):", min_value=1, value=15, key=f"cap_carr_{ik}")
                detalhes_arma = f" Acompanha {qtd_carr} carregador(es), com capacidade para {cap_carr} munições."
            elif tem_carr == "Não":
                detalhes_arma = " Não acompanha carregador."
        
        elif t_sel == "REVÓLVER":
            cap_tambor = st.number_input("Capacidade do tambor (munições):", min_value=1, max_value=12, value=6, key=f"cap_tambor_{ik}")
            tipo_abertura = st.selectbox("Abertura do tambor:", ["", "deslocamento lateral", "basculamento do cano", "janela lateral direita", "outra"], key=f"abertura_{ik}")
            abertura_txt = f" através de {tipo_abertura}" if tipo_abertura else ""
            detalhes_arma = f" Capacidade para {cap_tambor} munições{abertura_txt}."
            
        elif t_sel in ["ESPINGARDA", "GARRUCHA"]:
            num_canos = st.number_input("Número de canos:", min_value=1, max_value=4, value=1, key=f"num_canos_{ik}")
            num_gatilhos = st.number_input("Número de gatilhos:", min_value=1, max_value=4, value=1, key=f"num_gatilhos_{ik}")
            detalhes_arma = f" Possui {num_canos} cano(s) e {num_gatilhos} gatilho(s)."
            
        c1, c2 = st.columns(2)
        with c1:
            cao_arma = st.selectbox("Cão:", ["Aparente", "Oculto", "Não se aplica"], key=f"cao_{ik}")
            comp_cano = st.text_input("Comprimento do cano (Ex: 102 mm):", key=f"comp_cano_{ik}")
        with c2:
            alma_arma = st.selectbox("Alma do cano:", ["Raiada", "Lisa"], key=f"alma_{ik}")
            if alma_arma == "Raiada":
                num_raias = st.number_input("Número de raias:", min_value=1, value=6, key=f"raias_{ik}")
                sentido_raias = st.selectbox("Sentido:", ["Dextrógiras", "Sinistrógiras"], key=f"sentido_{ik}")
                alma_desc = f"raiada, com {num_raias} raias {sentido_raias.lower()}"
            else: alma_desc = "lisa"
        
        st.write("**Numeração e Identificação:**")
        num_status = st.selectbox("Situação da Numeração:", ["Íntegra", "Não aparente", "Suprimida", "Parcialmente visível", "Ausente"], key=f"num_status_{ik}")
        num_arma = ""
        metalo_txt = ""
        
        if num_status == "Íntegra": num_arma = st.text_input("Numeração Lida:", key=f"num_lida_{ik}")
        elif num_status != "Ausente":
            tipo_sup = st.selectbox("Método de supressão provável:", ["Por abrasão (lixada)", "Por punção (pinada)", "Indeterminado", "Outro"], key=f"tipo_sup_{ik}")
            if tipo_sup == "Outro": tipo_sup = st.text_input("Especifique a supressão:", key=f"sup_esp_{ik}")
            
            st.write("Exame Metalográfico (Reativo de Fry):")
            fry = st.selectbox("Resultado do Fry:", ["Não realizado", "Realizado - Revelou numeração", "Realizado - Revelou parcial", "Realizado - Negativo"], key=f"fry_{ik}")
            if "Revelou" in fry:
                num_revelada = st.text_input("Numeração Revelada:", key=f"num_rev_{ik}")
                metalo_txt = f"A arma foi submetida a exame metalográfico químico (Reativo de Fry), o qual revelou a numeração {num_revelada}."
            elif "Negativo" in fry:
                metalo_txt = "A arma foi submetida a exame metalográfico químico (Reativo de Fry), com resultado negativo para revelação."
                
        st.write("**Exames Finais:**")
        eficaz_arma = st.selectbox("Eficácia:", ["", "Eficaz para efetuar disparos.", "Ineficaz para efetuar disparos."], key=f"eficaz_arma_{ik}")
        residuografico = st.selectbox("Teste Residuográfico:", ["", "Negativo para disparo recente.", "Positivo para disparo recente.", "Não realizado."], key=f"resid_{ik}")
        lacre_saida_arma = st.text_input("Nº Lacre de Saída (Devolução da Arma):", key=f"lacre_saida_arma_{ik}")
        
        if st.button("Adicionar Arma"):
            desc_fisica = f"Cão {cao_arma.lower()}, cano com {comp_cano}, alma {alma_desc}.{detalhes_arma}"
            desc_num = num_arma if num_status == "Íntegra" else f"{num_status} ({tipo_sup})." if num_status != "Ausente" else "Ausente"
            st.session_state['itens_balistica'].append({
                "lacre": lacre_atual, "lacre_saida": lacre_saida_arma, "categoria": "Arma de Fogo", "tipo": tipo_arma, "fabricante": fab_arma, 
                "calibre": cal_arma, "estado": estado_arma, "caracteristicas": desc_fisica, "numeracao": desc_num, 
                "metalo": metalo_txt, "eficacia": eficaz_arma, "residuografico": residuografico
            })
            st.success("Arma adicionada!")

    # --- MUNIÇÕES ---
    elif tipo_item == "Munições":
        qtde_mun = st.number_input("Quantidade de Munições:", min_value=1, value=1, key=f"qtd_mun_{ik}")
        cal_mun = st.text_input("Calibre:", key=f"cal_mun_{ik}")
        marca_mun = st.text_input("Marca (Ex: CBC):", key=f"marca_mun_{ik}")
        lote_mun = st.text_input("Número do Lote (Opcional):", key=f"lote_mun_{ik}")
        desc_mun = st.text_input("Descrição dos projéteis (Ex: 10 ponta oca e 4 ponta plana):", key=f"desc_mun_{ik}")
        efic_mun = st.selectbox("Eficácia:", ["",
            "Houve êxito em picotar e deflagar. As munições testadas foram descartadas como resíduo balístico (estojos ejetados).", 
            "Falha na deflagração.", "Não testadas."], key=f"efic_mun_{ik}")
        lacre_saida_mun = st.text_input("Nº Lacre de Saída (Munições Restantes):", key=f"lacre_saida_mun_{ik}")
        
        if st.button("Adicionar Munições"):
            st.session_state['itens_balistica'].append({
                "lacre": lacre_atual, "lacre_saida": lacre_saida_mun, "categoria": "Munições", "quantidade": qtde_mun, 
                "calibre": cal_mun, "marca": marca_mun, "lote": lote_mun, "descricao": desc_mun, "eficacia": efic_mun
            })
            st.success("Munições adicionadas!")

    # --- ESTOJOS ---
    elif tipo_item == "Estojos":
        qtde_est = st.number_input("Quantidade de Estojos:", min_value=1, value=1, key=f"qtd_est_{ik}")
        cal_est = st.text_input("Calibre inscrito na base:", key=f"cal_est_{ik}")
        marca_est = st.text_input("Marca (Ex: CBC):", key=f"marca_est_{ik}")
        lote_est = st.text_input("Número do Lote (Opcional):", key=f"lote_est_{ik}")
        lacre_saida_est = st.text_input("Nº Lacre de Saída (Devolução dos Estojos):", key=f"lacre_saida_est_{ik}")
        
        if st.button("Adicionar Estojos"):
            st.session_state['itens_balistica'].append({
                "lacre": lacre_atual, "lacre_saida": lacre_saida_est, "categoria": "Estojos", "quantidade": qtde_est, 
                "calibre": cal_est, "marca": marca_est, "lote": lote_est
            })
            st.success("Estojos adicionados!")

    # --- PROJÉTIL ---
    elif tipo_item == "Projétil":
        fab_proj = st.text_input("Fabricante (Padrão: Não consta):", value="Não consta", key=f"fab_proj_{ik}")
        tipo_proj = st.text_input("Tipo (Ex: encamisado, ponta ogival):", key=f"tipo_proj_{ik}")
        massa_proj = st.number_input("Massa em gramas (g):", min_value=0.0, step=0.1, key=f"massa_proj_{ik}")
        diam_proj = st.number_input("Diâmetro em milímetros (mm):", min_value=0.0, step=0.1, key=f"diam_proj_{ik}")
        lacre_saida_proj = st.text_input("Nº Lacre de Saída (Devolução do Projétil):", key=f"lacre_saida_proj_{ik}")
        
        calibre_estimado = estimar_calibre(massa_proj, diam_proj) if massa_proj > 0 else "Aguardando dados..."
        st.info(f"🎯 **Calibre Estimado:** {calibre_estimado}")
        
        if st.button("Adicionar Projétil"):
            st.session_state['itens_balistica'].append({
                "lacre": lacre_atual, "lacre_saida": lacre_saida_proj, "categoria": "Projétil", "fabricante": fab_proj, 
                "tipo": tipo_proj, "massa": massa_proj, "diametro": diam_proj, "estimativa": calibre_estimado
            })
            st.success("Projétil adicionado!")

    # Botões de gerenciamento da lista e do formulário
    st.markdown("---")
    colB1, colB2 = st.columns(2)
    with colB1:
        if st.button("🆕 Novo Item (Limpar Formulário)", use_container_width=True):
            st.session_state['item_mk'] += 1
            st.rerun()
    with colB2:
        if st.button("🗑️ Remover Itens Adicionados", use_container_width=True):
            st.session_state['itens_balistica'] = []
            st.rerun()

# ==========================================
# 4. FOTOGRAFIAS 
# ==========================================
st.markdown("---")
st.header("4. Fotografias")

fotos_up = st.file_uploader("Arraste ou selecione as fotos aqui:", type=["jpg", "jpeg", "png"], accept_multiple_files=True, key=f"up_{st.session_state['uploader_key']}")

if fotos_up:
    nomes_atuais = [foto["nome"] for foto in st.session_state['fotos']]
    for f in fotos_up:
        if f.name not in nomes_atuais and f.name not in st.session_state['fotos_removidas']:
            img = Image.open(f)
            img = ImageOps.exif_transpose(img)
            st.session_state['fotos'].append({"imagem": img, "legenda": "Visão geral do item.", "nome": f.name})
    
    st.session_state['uploader_key'] += 1
    st.rerun()

if st.session_state['fotos']:
    st.write("### Fotos Adicionadas:")
    for i, foto in enumerate(st.session_state['fotos']):
        cf1, cf2, cf3 = st.columns([1, 3, 1])
        with cf1: st.image(foto["imagem"], width=100)
        with cf2: 
            nova_leg = st.text_input("Legenda:", value=foto["legenda"], key=f"leg_{i}_{mk}")
            st.session_state['fotos'][i]["legenda"] = nova_leg
        with cf3:
            if st.button("Remover", key=f"rem_{i}_{mk}"):
                st.session_state['fotos_removidas'].append(st.session_state['fotos'][i]["nome"])
                st.session_state['fotos'].pop(i)
                st.rerun()

# ==========================================
# 5. EDIÇÃO FINAL DO TEXTO
# ==========================================
st.markdown("---")
st.header("5. Revisão e Edição de Texto")
st.info("💡 **Dica:** O texto gerado automaticamente aparece abaixo. Você pode reescrever, corrigir erros ou adicionar informações. Use **asteriscos duplos** para deixar uma palavra em negrito no Word (Ex: **Tipo:**).")

texto_obj_default = ", ".join(obj_padrao)
if obj_complemento: texto_obj_default += f", {obj_complemento}"
if texto_obj_default: texto_obj_default += "."

texto_exames_gerado = ""
if len(st.session_state['itens_balistica']) > 0:
    lacres_dict = {}
    for item in st.session_state['itens_balistica']:
        l = item.get('lacre', '')
        if l not in lacres_dict: lacres_dict[l] = []
        lacres_dict[l].append(item)

    contador_lacre = 1
    for lacre, lista_itens in lacres_dict.items():
        texto_exames_gerado += f"{contador_lacre}. Foi recebido para exames, em saco plástico transparente, com lacre de entrada nº {lacre}, os seguintes itens:\n"
        for item in lista_itens:
            if item.get('categoria') == "Arma de Fogo":
                texto_exames_gerado += f"• **Tipo:** Arma de fogo, {item.get('tipo', '')}.\n"
                texto_exames_gerado += f"• **Fabricante/Modelo:** {item.get('fabricante', '')}.\n"
                texto_exames_gerado += f"• **Calibre:** {item.get('calibre', '')}.\n"
                texto_exames_gerado += f"• **Estado de Conservação:** {item.get('estado', '')}.\n"
                texto_exames_gerado += f"• **Características Físicas:** {item.get('caracteristicas', '')}\n"
                texto_exames_gerado += f"• **Numeração:** {item.get('numeracao', '')}\n"
                if item.get('metalo'): texto_exames_gerado += f"• **Exame Metalográfico:** {item.get('metalo', '')}\n"
                if item.get('residuografico'): texto_exames_gerado += f"• **Residuográfico:** {item.get('residuografico', '')}\n"
                if item.get('eficacia'): texto_exames_gerado += f"• **Eficácia:** {item.get('eficacia', '')}\n"
                if item.get('lacre_saida'): texto_exames_gerado += f"A arma foi acondicionada no lacre de saída nº {item.get('lacre_saida', '')}.\n"
            
            elif item.get('categoria') == "Munições":
                txt_lote = f", lote {item.get('lote')}" if item.get('lote') else ""
                texto_exames_gerado += f"• {item.get('quantidade', '')} munições intactas, calibre {item.get('calibre', '')}, marca {item.get('marca', '')}{txt_lote}, sendo {item.get('descricao', '')}.\n"
                if item.get('eficacia'): texto_exames_gerado += f"• **Eficácia:** As munições foram submetidas a testes mecânicos. {item.get('eficacia', '')}\n"
                if item.get('lacre_saida'): texto_exames_gerado += f"As munições restantes foram acondicionadas no lacre de saída nº {item.get('lacre_saida', '')}.\n"

            elif item.get('categoria') == "Estojos":
                txt_lote = f", lote {item.get('lote')}" if item.get('lote') else ""
                texto_exames_gerado += f"• {item.get('quantidade', '')} estojos deflagrados, calibre {item.get('calibre', '')}, marca {item.get('marca', '')}{txt_lote}.\n"
                if item.get('lacre_saida'): texto_exames_gerado += f"Os estojos foram acondicionados no lacre de saída nº {item.get('lacre_saida', '')}.\n"
            
            elif item.get('categoria') == "Projétil":
                texto_exames_gerado += f"• **Tipo:** Projétil.\n"
                texto_exames_gerado += f"• **Fabricante:** {item.get('fabricante', '')}.\n"
                texto_exames_gerado += f"• **Características Físicas:** Projétil {item.get('tipo', '')}, massa {item.get('massa', '')}g, diâmetro {item.get('diametro', '')} mm.\n"
                texto_exames_gerado += f"• **Calibre estimado:** {item.get('estimativa', '')} (com base na tabela pdf de especificações da CBC).\n"
                if item.get('lacre_saida'): texto_exames_gerado += f"O projétil foi acondicionado no lacre de saída nº {item.get('lacre_saida', '')}.\n"
            texto_exames_gerado += "\n"
        contador_lacre += 1

objetivo_final = st.text_area("Objetivo da Perícia (Editável):", value=texto_obj_default, height=68)
exames_final = st.text_area("Corpo dos Exames (Editável):", value=texto_exames_gerado, height=400)


# --- GERAÇÃO DO DOCUMENTO WORD ---
st.markdown("---")
if st.button("Criar Laudo (.docx)", type="primary", use_container_width=True):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Courier New'
    style.font.size = Pt(11)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    section = doc.sections[0]
    header = section.header
    for p in header.paragraphs: p.text = ""
    table = header.add_table(rows=1, cols=3, width=Cm(15.5))
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = table.columns[2].width = Cm(2.2); table.columns[1].width = Cm(11.1)
    for cell in table.columns[0].cells: cell.width = Cm(2.2)
    for cell in table.columns[1].cells: cell.width = Cm(11.1)
    for cell in table.columns[2].cells: cell.width = Cm(2.2)

    p_left = table.cell(0, 0).paragraphs[0]; p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists("logo_ssp.png"): p_left.add_run().add_picture("logo_ssp.png", width=Cm(1.8))
    
    p_center = table.cell(0, 1).paragraphs[0]; p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_h1 = p_center.add_run("SECRETARIA DA SEGURANÇA PÚBLICA\nSUPERINTENDÊNCIA DA POLÍCIA TÉCNICO-CIENTÍFICA\n")
    run_h1.font.size = Pt(11)
    run_h2 = p_center.add_run("INSTITUTO DE CRIMINALÍSTICA\n“PERITO CRIMINAL DR. OCTÁVIO EDUARDO DE BRITO ALVARENGA”\nNÚCLEO DE PERÍCIAS CRIMINALÍSTICAS DE SÃO JOSÉ DOS CAMPOS\nEQUIPE DE PERÍCIAS CRIMINALÍSTICAS DE GUARATINGUETÁ")
    run_h2.font.size = Pt(8)

    p_right = table.cell(0, 2).paragraphs[0]; p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists("logo_ic.png"): p_right.add_run().add_picture("logo_ic.png", width=Cm(1.8))

    if rep_input or bo_input:
        p_id = doc.add_paragraph()
        texto_id = []
        if rep_input: texto_id.append(f"REP {rep_input.upper()} / {rep_ano}")
        if bo_input: texto_id.append(f"BO {bo_input.upper()} / {bo_ano}")
        
        delegacia_txt = f" - {delegacia_selecionada}" if delegacia_selecionada else ""
        p_id.add_run(" - ".join(texto_id) + delegacia_txt)
        p_id.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # CORPO
    p_nat = doc.add_paragraph()
    # 1. NATUREZA ATUALIZADA AQUI:
    run = p_nat.add_run("1 – NATUREZA: Exame em Arma de Fogo e Munições"); run.bold = True; run.font.size = Pt(14)
    adicionar_borda_inferior(p_nat)
    
    del_nome = delegado_selecionado if delegado_selecionado else "__________"
    del_dp = delegacia_selecionada if delegacia_selecionada else "__________"
    preambulo = (f"Aos {data_extenso}, no Instituto de Criminalística, da Superintendência da Polícia Técnico-Científica, "
                 f"da Secretaria da Segurança Pública do Estado de São Paulo, de conformidade com o disposto no artigo 178 "
                 f"do Decreto-Lei nº. 3689, de 03 de outubro de 1941, pelo Diretor do Instituto de Criminalística, Ricardo Lopes Ortega, "
                 f"foi designado o Perito Criminal {perito_selecionado}, para proceder ao exame supracitado, em atendimento à requisição "
                 f"da Autoridade Policial, Dr(a). {del_nome}, titular/em exercício na {del_dp}.")
    doc.add_paragraph(preambulo)

    p_obj = doc.add_paragraph()
    run_obj = p_obj.add_run("2 - OBJETIVO DA PERÍCIA:"); run_obj.bold = True; run_obj.font.size = Pt(14)
    adicionar_borda_inferior(p_obj)
    doc.add_paragraph(objetivo_final)

    p_ex = doc.add_paragraph()
    run_ex = p_ex.add_run("3 – DOS EXAMES:"); run_ex.bold = True; run_ex.font.size = Pt(14)
    adicionar_borda_inferior(p_ex)

    for linha in exames_final.split('\n'):
        if not linha.strip():
            doc.add_paragraph()
            continue
            
        if linha.startswith("• "):
            p = doc.add_paragraph(style='List Bullet')
            linha_limpa = linha[2:]
        else:
            p = doc.add_paragraph()
            linha_limpa = linha
            
        partes = re.split(r'(\*\*.*?\*\*)', linha_limpa)
        for parte in partes:
            if parte.startswith('**') and parte.endswith('**'):
                p.add_run(parte[2:-2]).bold = True
            else:
                p.add_run(parte)

    # ILUSTRATIVO FOTOGRÁFICO
    if st.session_state['fotos']:
        # 2. QUEBRA DE PÁGINA ANTES DO FOTOGRÁFICO:
        doc.add_page_break()
        p_foto_cabecalho = doc.add_paragraph()
        run_foto = p_foto_cabecalho.add_run("4 – DO ILUSTRATIVO FOTOGRÁFICO:"); run_foto.bold = True; run_foto.font.size = Pt(14)
        adicionar_borda_inferior(p_foto_cabecalho)
        
        for foto in st.session_state['fotos']:
            img_stream = io.BytesIO()
            foto['imagem'].save(img_stream, format='PNG')
            img_stream.seek(0)
            
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = p_img.add_run()
            run_img.add_picture(img_stream, width=Cm(14))
            
            p_leg = doc.add_paragraph(f"Figura: {foto['legenda']}")
            p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ENCERRAMENTO PADRÃO
    p_relatar = doc.add_paragraph("Era o que havia a relatar.")
    p_relatar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_paginas = doc.add_paragraph("Este laudo vai impresso em ")
    adicionar_campo_numpages(p_paginas)  
    p_paginas.add_run(" páginas, além da capa, ficando arquivada cópia digital no sistema GDL da SPTC.")
    p_paginas.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_assinatura = doc.add_paragraph(); p_assinatura.paragraph_format.space_after = p_assinatura.paragraph_format.space_before = Pt(0)
    p_assinatura.add_run(perito_selecionado).bold = True
    p_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_cargo = doc.add_paragraph("Perito Criminal Relator"); p_cargo.paragraph_format.space_after = p_cargo.paragraph_format.space_before = Pt(0)
    p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf_doc = io.BytesIO(); doc.save(buf_doc); buf_doc.seek(0)
    
    if rep_input:
        nome_arquivo = f"Laudo_Balistica_REP_{rep_input}_{rep_ano}.docx"
    elif bo_input:
        nome_arquivo = f"Laudo_Balistica_BO_{bo_input}_{bo_ano}.docx"
    else:
        nome_arquivo = "Laudo_Balistica_Sem_BO_REP.docx"
        
    st.download_button("⬇️ Descarregar Laudo Final", buf_doc, nome_arquivo, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)