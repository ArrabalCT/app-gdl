import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import io
import os
from PIL import Image, ImageOps

# --- Funções de Formatação do Word ---
def adicionar_borda_inferior(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def adicionar_campo_numpages(paragraph):
    p = paragraph._p
    r1 = OxmlElement('w:r')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    r1.append(fldChar1)
    p.append(r1)
    r2 = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' NUMPAGES '
    r2.append(instrText)
    p.append(r2)
    r3 = OxmlElement('w:r')
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    r3.append(fldChar2)
    p.append(r3)
    r4 = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = 'xx'
    r4.append(t)
    p.append(r4)
    r5 = OxmlElement('w:r')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r5.append(fldChar3)
    p.append(r5)

# --- Variáveis de Sessão ---
if 'fotos' not in st.session_state: st.session_state['fotos'] = []
if 'mk' not in st.session_state: st.session_state['mk'] = 0 

mk = st.session_state['mk']

# --- Listas de Dados ---
delegados = ["Selecione..."] + sorted(["Adilson Antonio Marcondes dos Santos", "Adriane Goncalves", "Anisio Galdioli", "Cesar Aparecido Vieira da Silva", "Cristiane Correa de Freitas", "Daniel Souza Baptista de Castro", "Ernani Ronaldo Giannico Braga", "Fabio Germano Figueiredo Cabett", "Flavia Maria Rocha Rollo", "Francisco Sannini Neto", "Hugo Parreiras de Macedo", "Jose Marcelo Silva Hial", "Leonardo da Costa Ferreira", "Marcelo Vieira Cavalcante", "Mario Celso Ribeiro Senne", "Paulo Roberto Gruschka Castilho", "Paulo Sergio Barbosa", "Pedro Rossati", "Sergio Lucas Adler Guedes de Oliveira", "Vania Idalira Z. de Oliveira"]) + ["Outro..."]
peritos = ["Alexandre Rabello de Oliveira", "Bruna Fernandes Nogueira", "Claude Thiago Arrabal", "Jéssica Pereira Gonçalves", "Júlia Soares Melo", "Luiz Fausto Prado Vasques", "Marcelo Mourão Dantas", "Márcio Steinmetz Soares", "Sarah Costa Teixeira", "Ruan Carvalho de Souza"]
cidades = ["Selecione...", "Aparecida", "Cachoeira Paulista", "Canas", "Cunha", "Guaratinguetá", "Lorena", "Piquete", "Potim", "Roseira", "Outra..."]

dps_por_cidade = {
    "Aparecida": ["DEL.POL.APARECIDA"],
    "Canas": ["DEL.POL.CANAS"],
    "Cachoeira Paulista": ["DEL.POL.CACHOEIRA PAULISTA"],
    "Cunha": ["DEL.POL.CUNHA"],
    "Guaratinguetá": ["01º D.P. GUARATINGUETA", "02º D.P. GUARATINGUETA", "03º D.P. GUARATINGUETA", "DEL.INV.GER. GUARATINGUETA", "DEL.SEC.GUARATINGUETA", "DEL.SEC.GUARATINGUETA PLANTÃO", "DISE- DEL.SEC.GUARATINGUETA"],
    "Lorena": ["01º D.P. LORENA", "02º D.P. LORENA", "DEL.POL.LORENA"],
    "Piquete": ["DEL.POL.PIQUETE"],
    "Potim": ["DEL.POL.POTIM"],
    "Roseira": ["DEL.POL.ROSEIRA"]
}

# --- INTERFACE ---
st.title("Gerador de Laudos - Celular")

st.header("1. Cabeçalho e Identificação")

colR1, colR2 = st.columns(2)
with colR1: rep_input = st.text_input("Número REP:", value="", key=f"rep_{mk}")
with colR2: rep_ano = st.text_input("Ano REP:", value="2026", key=f"rep_ano_{mk}")

colBO1, colBO2 = st.columns(2)
with colBO1: bo_input = st.text_input("Número do BO:", value="", key=f"bo_{mk}").upper()
with colBO2: bo_ano = st.text_input("Ano do BO:", value="2026", key=f"ano_{mk}")

data_selecionada = st.date_input("Data do Laudo:", format="DD/MM/YYYY", key=f"data_{mk}")
meses = {1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril', 5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto', 9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'}
data_extenso = f"{data_selecionada.day} de {meses[data_selecionada.month]} de {data_selecionada.year}"

# Objetivo com complemento
objetivos_selecionados = st.multiselect("Objetivo da Perícia:", ["Extração de Dados - Cellebrite", "Constatação de danos", "Fotografação"], default=["Extração de Dados - Cellebrite", "Constatação de danos"], key=f"obj_{mk}")
objetivo_extra = st.text_input("Complemento do Objetivo (caso necessário):", key=f"obj_extra_{mk}")

perito_selecionado = st.selectbox("Perito Criminal:", peritos, index=peritos.index("Claude Thiago Arrabal"), key=f"per_{mk}")

# Autoridade Sem Pré-seleção
del_sel = st.selectbox("Autoridade Policial:", delegados, index=0, key=f"del_sel_{mk}")
if del_sel == "Outro...":
    delegado_selecionado = st.text_input("Digite o nome da Autoridade:", key=f"del_dig_{mk}")
elif del_sel == "Selecione...":
    delegado_selecionado = "[AUTORIDADE NÃO INFORMADA]"
else:
    delegado_selecionado = del_sel

# Cidade e DP Sem Pré-seleção
colC1, colC2 = st.columns(2)
with colC1: 
    cid_sel = st.selectbox("Cidade:", cidades, index=0, key=f"cid_sel_{mk}")
with colC2:
    if cid_sel == "Selecione...":
        dp_sel = st.selectbox("Delegacia:", ["Selecione a cidade primeiro..."], disabled=True, key=f"dp_sel_disabled_{mk}")
        delegacia_selecionada = "[DELEGACIA NÃO INFORMADA]"
    elif cid_sel in dps_por_cidade:
        opcoes_dp = ["Selecione..."] + dps_por_cidade[cid_sel] + ["Outra..."]
        dp_sel = st.selectbox("Delegacia:", opcoes_dp, index=0, key=f"dp_sel_{mk}")
    else:
        dp_sel = st.text_input("Digite a Delegacia:", key=f"dp_dig_{mk}")

if cid_sel != "Selecione...":
    if dp_sel == "Outra...":
        delegacia_selecionada = st.text_input("Nome da DP Específica:", key=f"dp_esp_{mk}")
    elif dp_sel == "Selecione...":
        delegacia_selecionada = "[DELEGACIA NÃO INFORMADA]"
    else:
        delegacia_selecionada = dp_sel

st.header("2. Cadeia de Custódia e Aparelhos")

# Lacre e Qtd Lado a Lado
colL1, colL2 = st.columns(2)
with colL1:
    lacre_saida = st.text_input("Lacre de Saída (Envio para o NI):", key=f"lacre_saida_{mk}")
with colL2:
    qtd_aparelhos = st.number_input("Quantidade de Aparelhos:", min_value=1, max_value=20, value=1, key=f"qtd_{mk}")

dados_aparelhos = []

for i in range(qtd_aparelhos):
    st.markdown(f"### 📱 Item {i+1}")
    lacre_ent = st.text_input(f"Lacre de Entrada Item {i+1}:", key=f"lacre_ent_{i}_{mk}")
    
    col1, col2 = st.columns(2)
    with col1:
        tipo_opcoes = ["Smartphone", "Tablet", "Feature Phone", "Outro"]
        tipo_sel = st.selectbox(f"Tipo (Item {i+1}):", tipo_opcoes, index=0, key=f"tipo_sel_{i}_{mk}")
        tipo_cel = st.text_input(f"Qual tipo? (Item {i+1})", key=f"tipo_dig_{i}_{mk}") if tipo_sel == "Outro" else tipo_sel

        marca_opcoes = ["Selecione..."] + sorted(["Motorola", "Samsung", "Apple", "Xiaomi", "LG", "Asus", "Nokia", "Positivo", "Realme", "Multilaser"]) + ["Outra"]
        marca_sel = st.selectbox(f"Marca (Item {i+1}):", marca_opcoes, index=0, key=f"marca_sel_{i}_{mk}")
        marca_cel = st.text_input(f"Qual marca? (Item {i+1})", key=f"marca_dig_{i}_{mk}") if marca_sel == "Outra" else marca_sel

        modelo_cel = st.text_input(f"Modelo (Item {i+1}):", key=f"mod_{i}_{mk}")

    with col2:
        cor_opcoes = ["Selecione..."] + sorted(["Preta", "Branca", "Prata", "Cinza", "Azul", "Dourada", "Rosa", "Vermelha", "Amarela", "Verde", "Roxa"]) + ["Outra"]
        cor_sel = st.selectbox(f"Cor (Item {i+1}):", cor_opcoes, index=0, key=f"cor_sel_{i}_{mk}")
        cor_cel = st.text_input(f"Qual cor? (Item {i+1})", key=f"cor_dig_{i}_{mk}") if cor_sel == "Outra" else cor_sel

        imei_cel = st.text_input(f"IMEI (Item {i+1}):", key=f"imei_{i}_{mk}")
        
        # Capa com "Não" como padrão
        tem_capa = st.radio(f"Capa? (Item {i+1})", ["Sim", "Não"], index=1, horizontal=True, key=f"capa_{i}_{mk}")

    st.markdown(f"**SIMCards - Item {i+1}**")
    sim_opcoes = sorted(["Vivo", "Tim", "Claro", "Oi", "Correios Cellular"]) + ["Outra", "Nenhum"]

    c_sim1, c_sim2 = st.columns(2)
    with c_sim1:
        s1_sel = st.selectbox(f"SIMCard 1 (Item {i+1}):", sim_opcoes, index=sim_opcoes.index("Nenhum"), key=f"s1_sel_{i}_{mk}")
        s1_txt = st.text_input(f"Qual op? (1)", key=f"s1_dig_{i}_{mk}") if s1_sel == "Outra" else s1_sel
        iccid1 = st.text_input(f"ICCID 1:", key=f"icc1_{i}_{mk}")
    with c_sim2:
        s2_sel = st.selectbox(f"SIMCard 2 (Item {i+1}):", sim_opcoes, index=sim_opcoes.index("Nenhum"), key=f"s2_sel_{i}_{mk}")
        s2_txt = st.text_input(f"Qual op? (2)", key=f"s2_dig_{i}_{mk}") if s2_sel == "Outra" else s2_sel
        iccid2 = st.text_input(f"ICCID 2:", key=f"icc2_{i}_{mk}")

    # Danos
    locais_opcoes = ["Tela (Display)", "Tampa traseira", "Lentes da câmera", "Botões", "Bordas/Laterais", "Película de proteção"]
    locais_selecionados = st.multiselect(f"Onde há danos? (Item {i+1}):", locais_opcoes, key=f"loc_{i}_{mk}")
    danos_detalhes = {}
    for local in locais_selecionados:
        cd1, cd2 = st.columns(2)
        with cd1: 
            t_dano = st.multiselect(f"Dano ({local}):", ["Fratura", "Quebra", "Trinco", "Riscos", "Atritamento"], key=f"td_{local}_{i}_{mk}")
        with cd2: 
            ext = st.text_input(f"Localização/Extensão ({local}):", key=f"ext_{local}_{i}_{mk}")
        danos_detalhes[local] = {"tipo": t_dano, "ext": ext}

    dados_aparelhos.append({
        "lacre": lacre_ent, "tipo": tipo_cel, "marca": marca_cel, "modelo": modelo_cel, 
        "cor": cor_cel, "imei": imei_cel, "capa": tem_capa, "s1": s1_txt, "icc1": iccid1,
        "s2": s2_txt, "icc2": iccid2, "danos": locais_selecionados, "detalhes": danos_detalhes
    })

# --- GERADOR DE TEXTO ---
txt_gerado = "Foi recebido para exame, lacrado(s) em sacos plásticos transparentes, os seguintes itens:\n\n"
for idx, ap in enumerate(dados_aparelhos):
    lacre_str = ap['lacre'] if ap['lacre'].strip() else "não informado"
    txt_gerado += f"Item {idx+1} --- Lacre de entrada {lacre_str}\n"
    
    mod_str = ap['modelo'].upper() if ap['modelo'] else "não aparente"
    marca_final = ap['marca'].upper() if ap['marca'] != "Selecione..." else "[MARCA NÃO INFORMADA]"
    cor_final = ap['cor'].lower() if ap['cor'] != "Selecione..." else "[COR NÃO INFORMADA]"
    
    capa_str = "acompanha capa de proteção" if ap['capa'] == "Sim" else "não acompanha capa de proteção"
        
    imei_str = f"IMEI {ap['imei']}" if ap['imei'].strip() else "IMEI não aparente"
    
    chips = []
    if ap['s1'] != "Nenhum": 
        chips.append(f"da operadora {ap['s1']}" + (f" (ICCID {ap['icc1']})" if ap['icc1'] else ""))
    if ap['s2'] != "Nenhum": 
        chips.append(f"da operadora {ap['s2']}" + (f" (ICCID {ap['icc2']})" if ap['icc2'] else ""))
    sim_str = "acompanhado de SIMCard(s) " + " e ".join(chips) if chips else "não acompanha SIMCard"

    txt_gerado += f"Trata-se de um aparelho do tipo {ap['tipo'].lower()}, marca {marca_final}, modelo {mod_str}, de cor predominante {cor_final}, que {capa_str}. O dispositivo apresenta {imei_str} e {sim_str}.\n"
    
    if not ap['danos']:
        txt_gerado += "Ao exame físico externo, o dispositivo não apresentava danos ou avarias visíveis em seu display ou carcaça.\n\n"
    else:
        linhas = [f"{l.lower()} com {', '.join(ap['detalhes'][l]['tipo']).lower()} ({ap['detalhes'][l]['ext'].lower()})" for l in ap['danos']]
        txt_gerado += f"Ao exame físico externo, constatou-se: {'; '.join(linhas)}.\n\n"

txt_gerado += f"O(s) aparelho(s) foram enviados para extração de dados no Núcleo de Informática de São José dos Campos em lacre {lacre_saida if lacre_saida else '[INSERIR LACRE DE SAÍDA]'}. O resultado seguirá em laudo complementar."

st.header("3. Edição e Fotos")
texto_final = st.text_area("Texto final:", value=txt_gerado, height=350, key=f"edit_{mk}")

# Upload de Fotos apenas
fotos_up = st.file_uploader("Carregar Fotos da Galeria:", type=['jpg', 'jpeg', 'png'], accept_multiple_files=True, key=f"up_{mk}")
if fotos_up:
    for f in fotos_up:
        if f.name not in [img.get('nome') for img in st.session_state['fotos']]:
            img = Image.open(io.BytesIO(f.getvalue()))
            st.session_state['fotos'].append({'img': ImageOps.exif_transpose(img), 'nome': f.name})

if st.session_state['fotos']:
    cols = st.columns(3)
    for i, fd in enumerate(st.session_state['fotos']):
        with cols[i % 3]:
            st.image(fd['img'], use_container_width=True)
            if st.button("❌", key=f"del_{i}"):
                st.session_state['fotos'].pop(i)
                st.rerun()

# --- FINALIZAÇÃO ---
st.header("4. Finalizar")
c1, c2 = st.columns(2)

with c1:
    if st.button("GERAR LAUDO (.DOCX)", type="primary", use_container_width=True):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Courier New'
        style.font.size = Pt(11)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Cabeçalho
        section = doc.sections[0]
        header = section.header
        for p in header.paragraphs: 
            p.text = ""
            
        table = header.add_table(rows=1, cols=3, width=Cm(15.5))
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        largura_lateral = Cm(2.2)
        largura_meio = Cm(11.1)

        table.columns[0].width = largura_lateral
        table.columns[1].width = largura_meio
        table.columns[2].width = largura_lateral
        
        for cell in table.columns[0].cells: cell.width = largura_lateral
        for cell in table.columns[1].cells: cell.width = largura_meio
        for cell in table.columns[2].cells: cell.width = largura_lateral
        
        p_left = table.cell(0, 0).paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if os.path.exists("logo_ssp.png"): 
            p_left.add_run().add_picture("logo_ssp.png", width=Cm(1.8))
            
        p_center = table.cell(0, 1).paragraphs[0]
        p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_h1 = p_center.add_run("SECRETARIA DA SEGURANÇA PÚBLICA\nSUPERINTENDÊNCIA DA POLÍCIA TÉCNICO-CIENTÍFICA\n")
        run_h1.font.size = Pt(11)
        run_h1.bold = False
        run_h2 = p_center.add_run("INSTITUTO DE CRIMINALÍSTICA\n“PERITO CRIMINAL DR. OCTÁVIO EDUARDO DE BRITO ALVARENGA”\nNÚCLEO DE PERÍCIAS CRIMINALÍSTICAS DE SÃO JOSÉ DOS CAMPOS\nEQUIPE DE PERÍCIAS CRIMINALÍSTICAS DE GUARATINGUETÁ")
        run_h2.font.size = Pt(8)
        run_h2.bold = False

        p_right = table.cell(0, 2).paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if os.path.exists("logo_ic.png"): 
            p_right.add_run().add_picture("logo_ic.png", width=Cm(1.8))

        if rep_input or bo_input:
            p_ident = doc.add_paragraph()
            p_ident.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ident_text = []
            if rep_input: ident_text.append(f"REP {rep_input} / {rep_ano}")
            if bo_input: ident_text.append(f"BO {bo_input} / {bo_ano} - {delegacia_selecionada}")
            p_ident.add_run(" | ".join(ident_text))
        
        # Corpo
        p_nat = doc.add_paragraph()
        run = p_nat.add_run("1 – NATUREZA: Peças")
        run.bold = True
        run.font.size = Pt(14)
        adicionar_borda_inferior(p_nat)
        
        preambulo = (f"Aos {data_extenso}, no Instituto de Criminalística, da Superintendência da Polícia Técnico-Científica, "
                     f"da Secretaria da Segurança Pública do Estado de São Paulo, de conformidade com o disposto no artigo 178 "
                     f"do Decreto-Lei nº. 3689, de 03 de outubro de 1941, pelo Diretor do Instituto de Criminalística, Ricardo Lopes Ortega, "
                     f"foi designado o Perito Criminal {perito_selecionado}, para proceder ao exame supracitado, em atendimento à requisição "
                     f"da Autoridade Policial, Dr(a). {delegado_selecionado}, titular/em exercício na {delegacia_selecionada}.")
        doc.add_paragraph(preambulo)

        p_obj = doc.add_paragraph()
        run_o = p_obj.add_run("2 - OBJETIVO DA PERÍCIA:")
        run_o.bold = True
        run_o.font.size = Pt(14)
        adicionar_borda_inferior(p_obj)
        
        objetivos_str = ", ".join(objetivos_selecionados) if objetivos_selecionados else "Não especificado"
        if objetivo_extra.strip():
            objetivos_str += f", {objetivo_extra.strip()}"
        doc.add_paragraph(f"Consta na requisição de exame: {objetivos_str}.")

        p_ex = doc.add_paragraph()
        run_e = p_ex.add_run("3 – DOS EXAMES:")
        run_e.bold = True
        run_e.font.size = Pt(14)
        adicionar_borda_inferior(p_ex)
        
        for linha in texto_final.split("\n"): 
            if linha.strip() == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(linha.strip())

        # Fotos
        if st.session_state['fotos']:
            doc.add_page_break()
            p_fotos = doc.add_paragraph()
            run_fotos = p_fotos.add_run("4 - REGISTRO FOTOGRÁFICO")
            run_fotos.bold = True
            run_fotos.font.size = Pt(14)
            adicionar_borda_inferior(p_fotos)
            
            for i, fd in enumerate(st.session_state['fotos']):
                img_io = io.BytesIO()
                fd['img'].convert('RGB').save(img_io, format='JPEG')
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(img_io, width=Cm(12))
                
                legenda = doc.add_paragraph(f"Fotografia {i+1}: Dispositivo examinado.")
                legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("")

        # Encerramento
        p_relatar = doc.add_paragraph("Era o que havia a relatar.")
        p_relatar.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_paginas = doc.add_paragraph("Este laudo vai impresso em ")
        adicionar_campo_numpages(p_paginas)  
        p_paginas.add_run(" páginas, além da capa, ficando arquivada cópia digital no sistema GDL da SPTC.")
        p_paginas.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        p_assinatura = doc.add_paragraph()
        p_assinatura.paragraph_format.space_after = Pt(0)
        p_assinatura.paragraph_format.space_before = Pt(0)
        p_assinatura.add_run(perito_selecionado).bold = True
        p_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_cargo = doc.add_paragraph("Perito Criminal Relator")
        p_cargo.paragraph_format.space_after = Pt(0)
        p_cargo.paragraph_format.space_before = Pt(0)
        p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        
        if rep_input:
            nome_arquivo = f"Laudo_Cel_REP_{rep_input}_{rep_ano}.docx"
        elif bo_input:
            nome_arquivo = f"Laudo_Cel_BO_{bo_input}_{bo_ano}.docx"
        else:
            nome_arquivo = "Laudo_Cel_Sem_BO_REP.docx"
            
        st.download_button("⬇️ DESCARREGAR LAUDO (.DOCX)", buf, nome_arquivo, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

with c2:
    if st.button("🔄 Novo(s) Celular(es) (Limpar Tudo)", type="secondary", use_container_width=True):
        current_mk = st.session_state.get('mk', 0)
        st.session_state.clear()
        st.session_state['mk'] = current_mk + 1
        st.rerun()